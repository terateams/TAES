# -*- coding: utf-8 -*-
from docx import Document

doc = Document()

# ===== 中文版 =====
doc.add_heading('3月活动网络与无线保障服务报价单（商务版）', level=0)
doc.add_paragraph('日期：2026-01-22')

doc.add_heading('1. 项目概述', level=1)
items = [
    ('项目名称', '3月活动网络与无线网络保障服务'),
    ('服务标的', '有线网络（宽带/海外访问保障）与无线网络（Wi-Fi）保障'),
    ('服务周期', '2026年3月（具体起止日期以双方确认的活动排期为准）'),
    ('交付目标', '确保活动期间网络可用、关键链路稳定，满足现场业务连续性要求'),
]
for k, v in items:
    p = doc.add_paragraph()
    p.add_run(k + '：').bold = True
    p.add_run(v)

doc.add_heading('2. 服务范围与交付口径', level=1)
doc.add_paragraph('本报价覆盖以下内容（按总包方式交付）：')
for s in ['基础宽带接入服务（300M）', '海外访问保障带宽（100M，按日本东京计价口径）', '无线网络保障（1个大区 + 3个小区，按 8 个 AP 测算）']:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('3. 报价明细（人民币）', level=1)
table = doc.add_table(rows=4, cols=6)
table.style = 'Table Grid'
hdr = table.rows[0].cells
headers = ['序号', '服务项', '规格/口径', '数量/测算', '单价', '小计']
for i, h in enumerate(headers):
    hdr[i].text = h
    hdr[i].paragraphs[0].runs[0].bold = True
data = [
    ('1', '基础宽带服务', '300M', '1项', '10,000 元', '10,000 元'),
    ('2', '海外访问保障', '100M（参考：日本东京 1M=300元）', '100M', '300 元/M', '30,000 元'),
    ('3', '无线保障', '1个大区 + 3个小区（按 8 个 AP 测算）', '8 个 AP', '2,500 元/AP', '20,000 元'),
]
for r, row_data in enumerate(data, 1):
    for c, val in enumerate(row_data):
        table.rows[r].cells[c].text = val

doc.add_heading('4. 费用合计', level=1)
p = doc.add_paragraph()
p.add_run('报价合计：').bold = True
p.add_run('60,000 元（6 万元）').bold = True

doc.add_heading('5. 商务条款与备注', level=1)
terms = [
    ('费用口径', '以上报价为打包价，包含本报价范围内的全部费用。'),
    ('配合事项', '实施过程中需要与酒店方（网络/弱电/机房/现场管理）配合对接，提供必要的施工窗口、机房/弱电间出入与现场协调支持。'),
    ('合规责任', '涉及的合规性（含流程与材料）由我们负责；酒店与客户配合提供必要授权与现场条件。'),
    ('范围变更', '如现场范围/点位/带宽/AP 数量等发生变化或新增需求，按变更内容另行评估与报价。'),
    ('报价有效期', '建议 15 天（或以双方书面确认的有效期为准）。'),
]
for k, v in terms:
    p = doc.add_paragraph()
    p.add_run(k + '：').bold = True
    p.add_run(v)

# ===== English Version =====
doc.add_page_break()
doc.add_heading('March Event Network & Wireless Services Quotation (Business Version)', level=0)
doc.add_paragraph('Date: 2026-01-22')

doc.add_heading('1. Project Overview', level=1)
items_en = [
    ('Project Name', 'March Event Network & Wireless Services'),
    ('Service Scope', 'Wired network (broadband / overseas access) and wireless network (Wi-Fi) support'),
    ('Service Period', 'March 2026 (exact dates subject to mutual confirmation based on event schedule)'),
    ('Delivery Objective', 'Ensure network availability and critical link stability during the event to meet on-site business continuity requirements'),
]
for k, v in items_en:
    p = doc.add_paragraph()
    p.add_run(k + ': ').bold = True
    p.add_run(v)

doc.add_heading('2. Scope of Services', level=1)
doc.add_paragraph('This quotation covers the following (delivered on a turnkey basis):')
for s in ['Basic broadband access service (300M)', 'Overseas access bandwidth (100M, priced at Tokyo, Japan reference rate)', 'Wireless network support (1 main zone + 3 sub-zones, based on 8 APs)']:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('3. Pricing Details (RMB)', level=1)
table2 = doc.add_table(rows=4, cols=6)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
headers_en = ['No.', 'Service Item', 'Specification', 'Qty / Basis', 'Unit Price', 'Subtotal']
for i, h in enumerate(headers_en):
    hdr2[i].text = h
    hdr2[i].paragraphs[0].runs[0].bold = True
data_en = [
    ('1', 'Basic Broadband Service', '300M', '1 item', '¥10,000', '¥10,000'),
    ('2', 'Overseas Access', '100M (Ref: Tokyo ¥300/M)', '100M', '¥300/M', '¥30,000'),
    ('3', 'Wireless Support', '1 main + 3 sub-zones (8 APs)', '8 APs', '¥2,500/AP', '¥20,000'),
]
for r, row_data in enumerate(data_en, 1):
    for c, val in enumerate(row_data):
        table2.rows[r].cells[c].text = val

doc.add_heading('4. Total', level=1)
p = doc.add_paragraph()
p.add_run('Quotation Total: ').bold = True
p.add_run('¥60,000 (RMB Sixty Thousand)').bold = True

doc.add_heading('5. Terms & Remarks', level=1)
terms_en = [
    ('Pricing Basis', 'The above is a package price and includes all costs within the scope of this quotation.'),
    ('Coordination Required', 'Implementation requires coordination with the hotel (network / low-voltage / server room / on-site management) to provide necessary construction windows, access to server/telecom rooms, and on-site support.'),
    ('Compliance Responsibility', 'All compliance matters (including procedures and documentation) are our responsibility; the hotel and client will provide necessary authorizations and on-site conditions.'),
    ('Scope Changes', 'Any changes to on-site scope, locations, bandwidth, or AP quantity will be evaluated and quoted separately.'),
    ('Quotation Validity', 'Recommended 15 days (or as mutually confirmed in writing).'),
]
for k, v in terms_en:
    p = doc.add_paragraph()
    p.add_run(k + ': ').bold = True
    p.add_run(v)

output_path = '/Users/yangjun/Workbench/TAES/Today/2026-01-22-网络与无线保障服务报价单.docx'
doc.save(output_path)
print('Done:', output_path)
